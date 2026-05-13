from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Dubs_Buddy_Setup_Guide.docx"
DUBS_IMAGE = ROOT / "references" / "style" / "Dubs.png"

INK = RGBColor(17, 24, 39)
MUTED = RGBColor(75, 85, 99)
PURPLE = RGBColor(75, 46, 131)
GOLD = RGBColor(183, 165, 122)
LIGHT_PURPLE = "F4F0FF"
LIGHT_GRAY = "F8FAFC"
BORDER = "D8DEE8"


def main():
    doc = Document()
    setup_document(doc)
    add_cover(doc)
    add_intro(doc)
    add_release_install(doc)
    doc.add_page_break()
    add_source_install(doc)
    add_codex_steps(doc)
    add_validation(doc)
    add_state_map(doc)
    add_dev_notes(doc)
    add_troubleshooting(doc)
    add_sharing_notes(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


def setup_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    set_style(styles["Normal"], "Arial", 10.5, INK)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.15
    set_style(styles["Title"], "Arial", 24, INK, bold=True)
    set_style(styles["Heading 1"], "Arial", 15, PURPLE, bold=True)
    set_style(styles["Heading 2"], "Arial", 12, PURPLE, bold=True)
    for style_name in ("Heading 1", "Heading 2"):
        styles[style_name].paragraph_format.space_before = Pt(14)
        styles[style_name].paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "Dubs Buddy Codex Pet"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(header.runs[0], size=9, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.text = "Unofficial fan project | v3.1.0"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(footer.runs[0], size=9, color=MUTED)


def add_cover(doc):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [Inches(5.0), Inches(1.45)])
    remove_table_borders(table)
    left, right = table.rows[0].cells

    p = left.paragraphs[0]
    run = p.add_run("Dubs Buddy Codex Pet")
    set_run(run, size=24, color=INK, bold=True)

    subtitle = left.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run("Setup Guide for GitHub and UW Sharing")
    set_run(run, size=13.5, color=MUTED)

    facts = [
        ("Version", "3.1.0"),
        ("Format", "Codex native custom pet"),
        ("Platforms", "macOS and Windows"),
        ("Repo", "github.com/FlalaGoGoGo/uw-dubs-codex-pet"),
    ]
    for label, value in facts:
        p = left.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        label_run = p.add_run(f"{label}: ")
        set_run(label_run, size=10.5, color=INK, bold=True)
        value_run = p.add_run(value)
        set_run(value_run, size=10.5, color=INK)

    if DUBS_IMAGE.exists():
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        right.paragraphs[0].add_run().add_picture(str(DUBS_IMAGE), width=Inches(1.25))
    shade_cell(right, LIGHT_PURPLE)
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    add_rule(doc, PURPLE)


def add_intro(doc):
    h1(doc, "1. What This Is")
    p(
        doc,
        "Dubs Buddy is an unofficial UW-inspired custom pet for Codex Desktop. It is not a separate desktop app; users install it through Codex's native Custom pets feature.",
    )
    callout(
        doc,
        "Codex currently expects a fixed 8 x 9 spritesheet: each frame is 192 x 208 px, and the final spritesheet is 1536 x 1872 px.",
    )


def add_release_install(doc):
    h1(doc, "2. Install From The Release Zip")
    numbered(
        doc,
        [
            "Download uw-dubs-codex-pet-v3.1.0.zip from the latest GitHub release.",
            "Unzip it. The extracted folder should be named dubs-buddy.",
            "Copy the dubs-buddy folder into your Codex custom pets directory.",
        ],
    )
    simple_table(
        doc,
        ["Platform", "Custom pets folder"],
        [
            ["macOS", "~/.codex/pets/dubs-buddy"],
            ["Windows", "%USERPROFILE%\\.codex\\pets\\dubs-buddy"],
        ],
        [1.3, 5.2],
    )
    h2(doc, "macOS")
    code(doc, "mkdir -p ~/.codex/pets\ncp -R dubs-buddy ~/.codex/pets/dubs-buddy")
    h2(doc, "Windows PowerShell")
    code(
        doc,
        '$petDir = "$env:USERPROFILE\\.codex\\pets"\nNew-Item -ItemType Directory -Force $petDir | Out-Null\nCopy-Item -Recurse -Force ".\\dubs-buddy" "$petDir\\dubs-buddy"',
    )


def add_source_install(doc):
    h1(doc, "3. Build From Source")
    p(doc, "Use this path if you want to inspect the build, regenerate previews, or contribute changes.")
    code(
        doc,
        "git clone https://github.com/FlalaGoGoGo/uw-dubs-codex-pet.git\ncd uw-dubs-codex-pet\nnpm install\nnpm run build\nnpm run verify\nnpm run install:pet",
    )
    simple_table(
        doc,
        ["Command", "Purpose"],
        [
            ["npm run build", "Generate dist/dubs-buddy, contact sheet, previews, galleries, and demo WebP."],
            ["npm run verify", "Validate manifest, spritesheet size, transparency, frame drift, and preview files."],
            ["npm run status:pet", "Compare the repo build with the installed Codex pet folder."],
            ["npm run install:pet", "Install the built pet into ~/.codex/pets/dubs-buddy."],
            ["npm run package:zip", "Create the versioned release zip."],
        ],
        [1.8, 4.7],
    )


def add_codex_steps(doc):
    h1(doc, "4. Enable In Codex")
    numbered(
        doc,
        [
            "Open Codex Desktop.",
            "Open Appearance.",
            "Go to Pets, then Custom pets.",
            "Click Refresh.",
            "Select Dubs Buddy.",
            "Click Wake Pet.",
        ],
    )
    callout(doc, "If Dubs Buddy does not appear, restart Codex and confirm the folder path exactly matches the install path.")


def add_validation(doc):
    h1(doc, "5. Validate The Install")
    simple_table(
        doc,
        ["Check", "Expected result"],
        [
            ["Folder", "dubs-buddy is inside the Codex pets directory."],
            ["Manifest", "pet.json has id dubs-buddy and spritesheetPath spritesheet.png."],
            ["Version", "VERSION.txt reports 3.1.0."],
            ["Spritesheet", "spritesheet.png is 1536 x 1872 with transparency."],
            ["Codex UI", "Dubs Buddy appears after Refresh and can be selected."],
        ],
        [1.5, 5.0],
    )
    code(doc, "npm run verify\nnpm run status:pet")


def add_state_map(doc):
    h1(doc, "6. Animation States")
    simple_table(
        doc,
        ["Row", "Codex state", "Dubs Buddy animation"],
        [
            ["0", "idle", "Sitting, breathing, blinking"],
            ["1", "running-right", "Right-facing side run"],
            ["2", "running-left", "Left-facing side run"],
            ["3", "waving", "Friendly paw wave"],
            ["4", "jumping", "Real jump / hover attention"],
            ["5", "failed", "Sad and crying"],
            ["6", "waiting", "Watching an hourglass"],
            ["7", "running", "Typing on a laptop"],
            ["8", "review", "Happy gift-box success"],
        ],
        [0.55, 1.65, 4.3],
    )


def add_dev_notes(doc):
    h1(doc, "7. Maintainer Notes")
    bullets(
        doc,
        [
            "Keep references/style/Dubs.png as the visual style reference.",
            "Place current 4 x 2, eight-frame reference sheets in references/source-sheets.",
            "Keep side-facing running rows no-logo to avoid reversed or doubled W marks.",
            "Run npm run prepare:release before publishing a GitHub release.",
            "Inspect dist/preview-contact-sheet.png, dist/gallery.html, dist/transition-gallery.html, and the demo WebP before publishing.",
        ],
    )


def add_troubleshooting(doc):
    h1(doc, "8. Troubleshooting")
    simple_table(
        doc,
        ["Symptom", "Likely fix"],
        [
            ["Dubs Buddy does not appear", "Check the install path and click Refresh in Custom pets."],
            ["Custom pet fails to load", "Run npm run verify and confirm spritesheet.png is exactly 1536 x 1872."],
            ["White box behind pet", "Rebuild; the generated spritesheet should have transparent corners."],
            ["Wrong version appears", "Run npm run status:pet, then reinstall with npm run install:pet."],
            ["Windows path confusion", "Use %USERPROFILE%\\.codex\\pets\\dubs-buddy."],
        ],
        [2.0, 4.5],
    )


def add_sharing_notes(doc):
    h1(doc, "9. Sharing Notes")
    p(
        doc,
        "This is a non-commercial, unofficial fan project. It is not affiliated with, sponsored by, or endorsed by the University of Washington or OpenAI.",
    )
    bullets(
        doc,
        [
            "UW Trademarks & Licensing: https://www.washington.edu/trademarks/",
            "UW Brand Trademarks & Licensing: https://www.washington.edu/brand/marketing-resources/trademarks-licensing/",
        ],
    )
    h2(doc, "LinkedIn Draft")
    callout(
        doc,
        "I built Dubs Buddy, an unofficial UW-inspired custom pet for Codex Desktop. Download the release zip, copy the dubs-buddy folder into Codex's custom pets directory, then enable it from Appearance -> Pets -> Custom pets. Go Dawgs, and may your tests pass.",
        fill="FFF8E1",
    )


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def p(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    set_run(run, size=10.5, color=INK)


def bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(3)
        para.paragraph_format.left_indent = Inches(0.32)
        para.paragraph_format.first_line_indent = Inches(-0.18)
        set_run(para.add_run(item), size=10.5, color=INK)


def numbered(doc, items):
    num_id = create_decimal_numbering(doc)
    for item in items:
        para = doc.add_paragraph()
        set_numbering(para, num_id)
        para.paragraph_format.space_after = Pt(3)
        para.paragraph_format.left_indent = Inches(0.32)
        para.paragraph_format.first_line_indent = Inches(-0.18)
        set_run(para.add_run(item), size=10.5, color=INK)


def code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [Inches(6.5)])
    cell = table.rows[0].cells[0]
    shade_cell(cell, LIGHT_GRAY)
    set_cell_border(cell, BORDER)
    set_cell_margins(cell, top=90, bottom=90, start=120, end=120)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(text.splitlines()):
        if idx:
            para.add_run().add_break()
        set_run(para.add_run(line), name="Courier New", size=9.2, color=INK)


def callout(doc, text, fill=LIGHT_PURPLE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [Inches(6.5)])
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    set_cell_border(cell, BORDER)
    set_cell_margins(cell, top=90, bottom=90, start=130, end=130)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    set_run(para.add_run(text), size=10.2, color=INK)


def simple_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table, [Inches(width) for width in widths])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, "EEE8F8")
        set_cell_border(cell, BORDER)
        set_cell_margins(cell, top=80, bottom=80, start=110, end=110)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_run(cell.paragraphs[0].add_run(header), size=10, color=PURPLE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_border(cell, BORDER)
            set_cell_margins(cell, top=80, bottom=80, start=110, end=110)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            set_run(cell.paragraphs[0].add_run(value), size=9.6, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_rule(doc, color):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(8)
    p_pr = para._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), rgb_to_hex(color))
    border.append(bottom)
    p_pr.append(border)


def create_decimal_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    for tag, value in (
        ("w:start", "1"),
        ("w:numFmt", "decimal"),
        ("w:lvlText", "%1."),
        ("w:lvlJc", "left"),
    ):
        node = OxmlElement(tag)
        node.set(qn("w:val"), value)
        lvl.append(node)

    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "460")
    ind.set(qn("w:hanging"), "260")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def set_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num = num_pr.find(qn("w:numId"))
    if num is None:
        num = OxmlElement("w:numId")
        num_pr.append(num)
    num.set(qn("w:val"), str(num_id))


def set_style(style, name, size, color, bold=False):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    style.element.rPr.rFonts.set(qn("w:ascii"), name)
    style.element.rPr.rFonts.set(qn("w:hAnsi"), name)


def set_run(run, name="Arial", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_table_width(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(int(width.inches * 1440) for width in widths)))
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width.inches * 1440)))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(widths[idx].inches * 1440)))


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, **margins):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in margins.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def remove_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                element = OxmlElement(f"w:{edge}")
                element.set(qn("w:val"), "nil")
                borders.append(element)
            tc_pr.append(borders)


def rgb_to_hex(color):
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


if __name__ == "__main__":
    main()
